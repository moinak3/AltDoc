from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "AltDoc_take_home_submission.md"
OUTPUT = ROOT / "AltDoc_take_home_submission.docx"


def set_paragraph_spacing(paragraph, *, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def set_run(run, *, size=11, bold=False, color="000000"):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Arial")
    font.set(qn("w:hAnsi"), "Arial")
    r_pr.append(font)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    r_pr.append(size)

    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(r_pr)
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_text_with_links(paragraph, text):
    url_pattern = re.compile(r"(https?://[^\s]+)")
    pos = 0
    for match in url_pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run(run)
        url = match.group(1)
        add_hyperlink(paragraph, url, url)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run(run)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    add_text_with_links(p, text)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_spacing(p, before=0, after=3)
        run = p.add_run(text)
        set_run(run, size=26, bold=False)
    elif level == 2:
        set_paragraph_spacing(p, before=20, after=6)
        run = p.add_run(text)
        set_run(run, size=20, bold=False)
    else:
        set_paragraph_spacing(p, before=18, after=6)
        run = p.add_run(text)
        set_run(run, size=16, bold=False)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4)
    add_text_with_links(p, text)
    return p


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15


def build():
    doc = Document()
    configure_doc(doc)

    lines = SOURCE.read_text().splitlines()
    pending = []

    def flush_pending():
        nonlocal pending
        if pending:
            add_paragraph(doc, " ".join(pending).strip())
            pending = []

    for raw in lines:
        line = raw.strip()
        if not line:
            flush_pending()
            continue
        if line.startswith("# "):
            flush_pending()
            add_heading(doc, line[2:].strip(), 1)
            continue
        if line.startswith("## "):
            flush_pending()
            add_heading(doc, line[3:].strip(), 2)
            continue
        if line.startswith("- "):
            flush_pending()
            add_bullet(doc, line[2:].strip())
            continue
        pending.append(line.replace("`", ""))

    flush_pending()
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
